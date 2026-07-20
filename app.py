# -*- coding: utf-8 -*-
"""
이노티움 회의록 자동 생성 — 클라우드 버전
- 로컬 PC 불필요, 어디서나 사용 가능
- 음성 파일 → Groq Whisper (무료) → GPT-4o → 회의록 DOCX
- 텍스트 / 문서 파일 → GPT-4o → 회의록 DOCX
- 레드마인 저장은 선택 버튼 (사내망 연결 시만 작동)
"""

import os, json, tempfile, logging
from datetime import datetime
from pathlib import Path

from flask import Flask, request, jsonify, send_file, render_template_string
from flask_cors import CORS
from dotenv import load_dotenv
from openai import AzureOpenAI
import requests as http_req

load_dotenv()
logging.basicConfig(level=logging.INFO)
log = logging.getLogger("meeting-cloud")

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024
CORS(app, resources={r"/*": {"origins": "*"}})

UPLOAD_TMP = Path(tempfile.gettempdir()) / "meeting_cloud"
UPLOAD_TMP.mkdir(exist_ok=True)

# ── Azure OpenAI (GPT-4o — 회의록 생성용) ────────────────────────────────────
AZURE_ENDPOINT   = os.environ.get("AZURE_OPENAI_ENDPOINT", "")
AZURE_KEY        = os.environ.get("AZURE_OPENAI_KEY", "")
AZURE_DEPLOY     = os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4o")
AZURE_API_VER    = os.environ.get("AZURE_OPENAI_API_VERSION", "2025-01-01-preview")

# ── Groq (Whisper 무료 STT) ───────────────────────────────────────────────────
GROQ_API_KEY     = os.environ.get("GROQ_API_KEY", "")
GROQ_WHISPER     = os.environ.get("GROQ_WHISPER_MODEL", "whisper-large-v3-turbo")

# ── fallback: OpenAI / Azure Whisper ─────────────────────────────────────────
OPENAI_API_KEY   = os.environ.get("OPENAI_API_KEY", "")
AZURE_WHISPER    = os.environ.get("AZURE_OPENAI_WHISPER_DEPLOYMENT", "")

# 레드마인 (선택 — 사내망 전용)
REDMINE_URL      = os.environ.get("REDMINE_URL", "")
REDMINE_API_KEY  = os.environ.get("REDMINE_API_KEY", "")
REDMINE_PROJECT  = os.environ.get("REDMINE_MEETING_PROJECT", "weekly_meeting_notes")
REDMINE_TRACKER_UPLOAD = int(os.environ.get("REDMINE_TRACKER_UPLOAD", "10"))

ALLOWED_AUDIO = {".m4a", ".mp3", ".wav", ".mp4", ".webm", ".ogg", ".flac"}
ALLOWED_DOCS  = {".txt", ".md", ".csv", ".pdf", ".docx", ".doc",
                 ".hwp", ".hwpx", ".rtf", ".xlsx", ".pptx"}

azure_client = AzureOpenAI(
    azure_endpoint=AZURE_ENDPOINT,
    api_key=AZURE_KEY,
    api_version=AZURE_API_VER,
) if AZURE_KEY else None


# ── HTML ─────────────────────────────────────────────────────────────────────
HTML = r"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>이노티움 회의록 자동 생성</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Malgun Gothic','맑은 고딕',Segoe UI,sans-serif;background:#f0f4f8;min-height:100vh;padding:24px 16px}
.wrap{max-width:860px;margin:0 auto}
header{text-align:center;padding:28px 0 22px}
header h1{font-size:1.85em;color:#1a365d;margin-bottom:6px}
header p{color:#718096;font-size:.9em}
.cloud-badge{display:inline-block;background:#ebf8ff;color:#2b6cb0;border:1px solid #bee3f8;border-radius:20px;padding:3px 12px;font-size:.78em;font-weight:600;margin-top:6px}
.card{background:#fff;border-radius:14px;padding:28px;margin-bottom:18px;box-shadow:0 2px 10px rgba(0,0,0,.07)}
.card h2{font-size:1.05em;color:#2d3748;margin-bottom:18px;padding-bottom:10px;border-bottom:2px solid #e2e8f0}
.tabs{display:flex;gap:8px;margin-bottom:18px}
.tab{flex:1;padding:10px;border:2px solid #e2e8f0;background:#fff;border-radius:9px;cursor:pointer;font-family:inherit;font-size:.93em;color:#718096;transition:.2s}
.tab.on{border-color:#4299e1;background:#ebf8ff;color:#2b6cb0;font-weight:600}
.pane{display:none}.pane.on{display:block}
.drop{border:2px dashed #cbd5e0;border-radius:10px;padding:40px 20px;text-align:center;cursor:pointer;transition:.2s;background:#f7fafc}
.drop:hover,.drop.over{border-color:#4299e1;background:#ebf8ff}
.drop .ico{font-size:2.8em;margin-bottom:8px}
.drop p{color:#718096;margin:4px 0}.drop .hint{font-size:.8em;color:#a0aec0;margin-top:6px}
#fname,#doc-fname{margin-top:8px;font-size:.88em;color:#4299e1;font-weight:600}
.notice{background:#fffbeb;border:1px solid #f6e05e;border-radius:8px;padding:10px 14px;font-size:.83em;color:#744210;margin-top:10px;line-height:1.5}
.notice-blue{background:#ebf8ff;border:1px solid #bee3f8;border-radius:8px;padding:10px 14px;font-size:.83em;color:#2c5282;margin-top:10px;line-height:1.5}
textarea{width:100%;min-height:200px;border:2px solid #e2e8f0;border-radius:8px;padding:12px;font-family:inherit;font-size:.9em;resize:vertical;outline:none;transition:.2s}
textarea:focus{border-color:#4299e1}
.grid{display:grid;grid-template-columns:1fr 1fr;gap:12px}
.full{grid-column:1/-1}
label{display:block;font-size:.83em;color:#4a5568;margin-bottom:4px;font-weight:500}
input[type=text],input[type=datetime-local]{width:100%;padding:9px 12px;border:2px solid #e2e8f0;border-radius:8px;font-family:inherit;font-size:.9em;outline:none;transition:.2s}
input:focus{border-color:#4299e1}
.btn-go{width:100%;padding:14px;background:linear-gradient(135deg,#4299e1,#3182ce);color:#fff;border:none;border-radius:10px;font-family:inherit;font-size:1.05em;font-weight:600;cursor:pointer;transition:.2s;margin-top:8px}
.btn-go:hover{background:linear-gradient(135deg,#3182ce,#2b6cb0);transform:translateY(-1px)}
.btn-go:disabled{background:#a0aec0;cursor:not-allowed;transform:none}
.prog{display:none}
.bar-wrap{background:#e2e8f0;border-radius:10px;height:8px;margin:12px 0;overflow:hidden}
.bar{height:100%;background:linear-gradient(90deg,#4299e1,#63b3ed);width:0%;border-radius:10px;transition:width .5s;animation:pulse 1.8s infinite}
@keyframes pulse{0%,100%{opacity:1}50%{opacity:.55}}
.stxt{color:#4a5568;font-size:.93em;text-align:center;margin-top:6px}
.result{display:none}
.preview{background:#f7fafc;border:1px solid #e2e8f0;border-radius:8px;padding:20px;margin-bottom:18px;white-space:pre-wrap;font-size:.87em;line-height:1.8;max-height:520px;overflow-y:auto;font-family:'Consolas','Courier New',monospace}
.btn-dl{width:100%;padding:13px;background:linear-gradient(135deg,#48bb78,#38a169);color:#fff;border:none;border-radius:10px;font-family:inherit;font-size:.98em;font-weight:600;cursor:pointer;transition:.2s}
.btn-dl:hover{background:linear-gradient(135deg,#38a169,#2f855a);transform:translateY(-1px)}
.btn-dl:disabled{background:#a0aec0;cursor:not-allowed;transform:none}
.btn-reset{width:100%;padding:10px;background:#fff;color:#718096;border:2px solid #e2e8f0;border-radius:10px;font-family:inherit;font-size:.88em;cursor:pointer;margin-top:8px;transition:.2s}
.btn-reset:hover{border-color:#a0aec0;color:#4a5568}
/* 레드마인 버튼 (선택) */
.rm-section{margin-top:12px;padding-top:12px;border-top:1px solid #e2e8f0}
.rm-section p{font-size:.8em;color:#a0aec0;margin-bottom:8px;text-align:center}
.btn-rm{width:100%;padding:11px;background:#fff;color:#805ad5;border:2px solid #e9d8fd;border-radius:10px;font-family:inherit;font-size:.92em;font-weight:600;cursor:pointer;transition:.2s}
.btn-rm:hover{background:#faf5ff;border-color:#b794f4}
.btn-rm:disabled{opacity:.5;cursor:not-allowed}
.rm-ok{background:#f0fff4;border:1px solid #9ae6b4;border-radius:8px;padding:10px 14px;color:#276749;font-size:.85em;margin-top:8px;display:none}
.rm-ok a{color:#276749;font-weight:700}
.rm-err{background:#fff5f5;border:1px solid #feb2b2;border-radius:8px;padding:10px 14px;color:#c53030;font-size:.85em;margin-top:8px;display:none}
.err{background:#fff5f5;border:1px solid #feb2b2;border-radius:8px;padding:12px 16px;color:#c53030;font-size:.88em;display:none;margin-top:10px;line-height:1.5}
/* 문서 업로드 */
#doc-status{display:none;margin-top:10px}
#doc-result{display:none;margin-top:10px;background:#f7fafc;border:1px solid #e2e8f0;border-radius:8px;padding:12px;font-size:.8em;color:#4a5568;max-height:140px;overflow-y:auto;white-space:pre-wrap}
@media(max-width:600px){.grid{grid-template-columns:1fr}.full{grid-column:1}}
</style>
</head>
<body>
<div class="wrap">
  <header>
    <h1>📋 회의록 자동 생성</h1>
    <p>이노티움(주) | 음성파일 · 문서 · 텍스트 → GPT-4o → 표준 회의록 DOCX</p>
    <span class="cloud-badge">☁️ Cloud · 어디서나 사용 가능</span>
  </header>

  <div class="card" id="icard">
    <h2>1단계 · 회의 내용 입력</h2>
    <div class="tabs">
      <button class="tab on" onclick="sw('audio',this)">🎤 음성파일</button>
      <button class="tab"    onclick="sw('doc',this)">📄 문서 파일</button>
      <button class="tab"    onclick="sw('text',this)">📝 텍스트 입력</button>
    </div>

    <div class="pane on" id="pane-audio">
      <div class="drop" id="drop" onclick="document.getElementById('fi').click()">
        <div class="ico">🎙️</div>
        <p><strong>클릭하거나 드래그하세요</strong></p>
        <p class="hint">m4a · mp3 · wav · mp4 · webm · ogg · flac &nbsp;|&nbsp; 최대 200MB</p>
        <div id="fname"></div>
      </div>
      <input type="file" id="fi" accept=".m4a,.mp3,.wav,.mp4,.webm,.ogg,.flac" style="display:none">
      <div class="notice-blue" id="stt-status">ⓘ 음성 처리 중…</div>
    </div>

    <div class="pane" id="pane-doc">
      <div class="drop" id="drop-doc" onclick="document.getElementById('fi-doc').click()">
        <div class="ico">📄</div>
        <p><strong>클릭하거나 드래그하세요</strong></p>
        <p class="hint">txt · pdf · docx · doc · hwp · hwpx · xlsx · pptx · md · csv · rtf &nbsp;|&nbsp; 제한 없음</p>
        <div id="doc-fname"></div>
      </div>
      <input type="file" id="fi-doc" accept=".txt,.pdf,.docx,.doc,.hwp,.hwpx,.xlsx,.pptx,.md,.csv,.rtf" style="display:none">
      <div id="doc-status">
        <div class="bar-wrap"><div class="bar" id="doc-bar" style="width:0%"></div></div>
        <p class="stxt" id="doc-stxt">텍스트 추출 중…</p>
      </div>
      <div id="doc-result"></div>
    </div>

    <div class="pane" id="pane-text">
      <textarea id="ti" placeholder="회의 녹취 내용이나 메모를 붙여넣으세요&#10;&#10;예) 김대표: 이번 분기 매출 현황을 보고해주세요.&#10;영업팀장: 목표 대비 87% 달성했습니다."></textarea>
    </div>
  </div>

  <div class="card">
    <h2>2단계 · 회의 정보 <span style="color:#a0aec0;font-size:.82em">(선택)</span></h2>
    <div class="grid">
      <div><label>회의 일시</label><input type="datetime-local" id="mdate"></div>
      <div><label>회의 장소</label><input type="text" id="mplace" placeholder="예) 지혜의 방"></div>
      <div class="full"><label>참석자</label><input type="text" id="matt" placeholder="예) 이노티움: 김종필 대표, 배진섭 과장"></div>
      <div class="full"><label>작성자</label><input type="text" id="mauth" placeholder="예) 배진섭"></div>
    </div>
  </div>

  <button class="btn-go" id="bgo" onclick="go()">✨ 회의록 생성하기</button>
  <div class="err" id="err"></div>

  <div class="card prog" id="prog">
    <h2 id="prog-title">처리 중…</h2>
    <div class="bar-wrap"><div class="bar" id="bar"></div></div>
    <p class="stxt" id="stxt">요청 전송 중…</p>
  </div>

  <div class="card result" id="res">
    <h2>✅ 회의록 완성</h2>
    <div class="preview" id="prev"></div>
    <button class="btn-dl" id="bdl" onclick="dl()">📄 DOCX 파일 다운로드</button>

    <!-- 레드마인 저장 — 사내망 연결 시만 작동 -->
    <div class="rm-section">
      <p>📋 레드마인에 저장 (사내망 연결 시만 가능)</p>
      <button class="btn-rm" id="brm" onclick="saveRedmine()">📋 레드마인에 저장하기</button>
      <div class="rm-ok" id="rm-ok"></div>
      <div class="rm-err" id="rm-err"></div>
    </div>

    <button class="btn-reset" onclick="resetAll()">↩ 새 회의록 작성</button>
  </div>
</div>

<script>
let tab='audio', file=null, gdata=null, docText=null;

function sw(t,btn){
  tab=t;
  document.querySelectorAll('.tab').forEach(b=>b.classList.remove('on'));
  btn.classList.add('on');
  document.querySelectorAll('.pane').forEach(p=>p.classList.remove('on'));
  document.getElementById('pane-'+t).classList.add('on');
}

// 음성 파일
const drop=document.getElementById('drop'), fi=document.getElementById('fi');
drop.addEventListener('dragover',e=>{e.preventDefault();drop.classList.add('over')});
drop.addEventListener('dragleave',()=>drop.classList.remove('over'));
drop.addEventListener('drop',e=>{e.preventDefault();drop.classList.remove('over');if(e.dataTransfer.files[0])setFile(e.dataTransfer.files[0])});
fi.addEventListener('change',e=>{if(e.target.files[0])setFile(e.target.files[0])});
function setFile(f){file=f;document.getElementById('fname').textContent='📎 '+f.name+' ('+(f.size/1048576).toFixed(1)+'MB)'}

// 문서 파일
const dropDoc=document.getElementById('drop-doc'), fiDoc=document.getElementById('fi-doc');
dropDoc.addEventListener('dragover',e=>{e.preventDefault();dropDoc.classList.add('over')});
dropDoc.addEventListener('dragleave',()=>dropDoc.classList.remove('over'));
dropDoc.addEventListener('drop',e=>{e.preventDefault();dropDoc.classList.remove('over');if(e.dataTransfer.files[0])extractDoc(e.dataTransfer.files[0])});
fiDoc.addEventListener('change',e=>{if(e.target.files[0])extractDoc(e.target.files[0])});

async function extractDoc(f){
  docText=null;
  document.getElementById('doc-fname').textContent='📎 '+f.name;
  document.getElementById('doc-status').style.display='block';
  document.getElementById('doc-result').style.display='none';
  document.getElementById('doc-bar').style.width='40%';
  const fd=new FormData(); fd.append('file',f);
  try{
    document.getElementById('doc-bar').style.width='70%';
    const r=await fetch('/extract-text',{method:'POST',body:fd});
    const d=await r.json();
    document.getElementById('doc-bar').style.width='100%';
    document.getElementById('doc-status').style.display='none';
    if(!r.ok||d.error){showErr(d.error||'추출 실패');return;}
    docText=d.text;
    const prev=document.getElementById('doc-result');
    prev.style.display='block';
    prev.textContent='✅ '+d.chars.toLocaleString()+'자 추출됨 ('+f.name+')\n\n'+d.text.slice(0,300)+(d.text.length>300?'\n...':'');
  }catch(e){showErr('추출 오류: '+e.message);document.getElementById('doc-status').style.display='none';}
}

function showErr(m){const e=document.getElementById('err');e.innerHTML='⚠️ '+m;e.style.display='block'}
function hideErr(){document.getElementById('err').style.display='none'}
function setBar(pct,txt){document.getElementById('bar').style.width=pct+'%';document.getElementById('stxt').textContent=txt}

async function go(){
  hideErr();
  const textVal=document.getElementById('ti').value.trim();
  if(tab==='audio'&&!file){showErr('음성 파일을 선택해주세요.');return}
  if(tab==='doc'&&!docText){showErr('문서 파일을 먼저 업로드해주세요.');return}
  if(tab==='text'&&!textVal){showErr('텍스트를 입력해주세요.');return}

  document.getElementById('bgo').disabled=true;
  document.getElementById('icard').style.display='none';
  document.getElementById('prog').style.display='block';
  setBar(5,'요청 전송 중…');

  const meta={
    date:document.getElementById('mdate').value,
    place:document.getElementById('mplace').value,
    attendees:document.getElementById('matt').value,
    author:document.getElementById('mauth').value,
  };

  let req;
  if(tab==='audio'){
    const fd=new FormData();
    fd.append('mode','audio'); fd.append('audio',file);
    Object.entries(meta).forEach(([k,v])=>fd.append(k,v));
    document.getElementById('prog-title').textContent='음성 변환 중… (시간이 걸릴 수 있어요)';
    setTimeout(()=>setBar(20,'음성 파일 업로드 중…'),500);
    setTimeout(()=>setBar(50,'STT 변환 중…'),5000);
    setTimeout(()=>setBar(80,'GPT-4o 회의록 작성 중…'),20000);
    req=fetch('/generate',{method:'POST',body:fd});
  } else {
    const inputText=tab==='doc'?docText:textVal;
    document.getElementById('prog-title').textContent='GPT-4o 회의록 작성 중…';
    setTimeout(()=>setBar(40,'AI 분석 중…'),500);
    req=fetch('/generate',{
      method:'POST',
      headers:{'Content-Type':'application/json; charset=utf-8'},
      body:JSON.stringify({mode:'text',text:inputText,...meta}),
    });
  }

  try{
    const r=await req;
    const d=await r.json();
    if(!r.ok||d.error)throw new Error(d.error||'서버 오류');
    setBar(100,'완료!');
    gdata=d;
    setTimeout(()=>{
      document.getElementById('prog').style.display='none';
      document.getElementById('res').style.display='block';
      document.getElementById('prev').textContent=d.preview;
    },600);
  }catch(e){
    document.getElementById('prog').style.display='none';
    document.getElementById('icard').style.display='block';
    document.getElementById('bgo').disabled=false;
    showErr(e.message);
  }
}

async function dl(){
  if(!gdata)return;
  const btn=document.getElementById('bdl');
  btn.textContent='준비 중…'; btn.disabled=true;
  try{
    const r=await fetch('/download',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(gdata)});
    if(r.ok){
      const blob=await r.blob();
      const a=document.createElement('a');
      a.href=URL.createObjectURL(blob);
      a.download=gdata.filename||'회의록.docx';
      a.click();
    }else{const d=await r.json();alert('다운로드 실패: '+(d.error||'오류'));}
  }finally{btn.textContent='📄 DOCX 파일 다운로드';btn.disabled=false;}
}

async function saveRedmine(){
  if(!gdata)return;
  const btn=document.getElementById('brm');
  btn.textContent='저장 중…'; btn.disabled=true;
  document.getElementById('rm-ok').style.display='none';
  document.getElementById('rm-err').style.display='none';
  try{
    const r=await fetch('/redmine',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(gdata)});
    const d=await r.json();
    if(r.ok&&d.url){
      const box=document.getElementById('rm-ok');
      box.innerHTML='✅ 저장 완료! <a href="'+d.url+'" target="_blank">🔗 레드마인에서 보기</a>';
      box.style.display='block';
      btn.textContent='✅ 저장됨';
    }else{throw new Error(d.error||'저장 실패');}
  }catch(e){
    const errBox=document.getElementById('rm-err');
    errBox.textContent='❌ '+e.message+(e.message.includes('fetch')?'\n(사내망 연결 필요)':'');
    errBox.style.display='block';
    btn.textContent='📋 레드마인에 저장하기';
    btn.disabled=false;
  }
}

function resetAll(){
  gdata=null;file=null;docText=null;
  document.getElementById('fname').textContent='';
  document.getElementById('fi').value='';
  document.getElementById('ti').value='';
  document.getElementById('doc-fname').textContent='';
  document.getElementById('fi-doc').value='';
  document.getElementById('doc-result').style.display='none';
  document.getElementById('doc-status').style.display='none';
  document.getElementById('res').style.display='none';
  document.getElementById('icard').style.display='block';
  document.getElementById('bgo').disabled=false;
  document.getElementById('brm').textContent='📋 레드마인에 저장하기';
  document.getElementById('brm').disabled=false;
  document.getElementById('rm-ok').style.display='none';
  document.getElementById('rm-err').style.display='none';
  hideErr();
  sw('audio',document.querySelectorAll('.tab')[0]);
}

(()=>{
  const n=new Date(); n.setMinutes(n.getMinutes()-n.getTimezoneOffset());
  document.getElementById('mdate').value=n.toISOString().slice(0,16);
  // STT 상태 표시
  fetch('/stt-status').then(r=>r.json()).then(d=>{
    const el=document.getElementById('stt-status');
    if(d.available){
      el.className='notice-blue';
      el.textContent='ⓘ 음성 변환: '+d.engine+' | 회의록 작성: GPT-4o';
    }else{
      el.className='notice';
      el.innerHTML='⚠️ 현재 음성 변환 기능이 비활성화되어 있습니다.<br>'
        +'📄 <strong>문서 파일</strong> 또는 📝 <strong>텍스트 입력</strong> 탭을 이용해주세요.';
    }
  }).catch(()=>{});
})();
</script>
</body>
</html>"""


# ── STT 엔진 감지 ────────────────────────────────────────────────────────────
def _stt_engine():
    """우선순위: Groq → OpenAI → Azure Whisper → 로컬 Whisper"""
    if GROQ_API_KEY:
        return "groq", f"Groq {GROQ_WHISPER} (무료)"
    if OPENAI_API_KEY:
        return "openai", "OpenAI Whisper API"
    if AZURE_WHISPER and azure_client:
        return "azure", f"Azure Whisper ({AZURE_WHISPER})"
    try:
        from faster_whisper import WhisperModel
        return "local", "로컬 Whisper (faster-whisper)"
    except ImportError:
        return None, None


# ── 음성 → 텍스트 ─────────────────────────────────────────────────────────────
def transcribe(file_path: str) -> str:
    engine, _ = _stt_engine()
    fname = Path(file_path).name

    if engine == "groq":
        from groq import Groq
        import os as _os
        file_size = _os.path.getsize(file_path)
        if file_size > 24 * 1024 * 1024:  # 24MB Groq 제한
            raise ValueError(
                f"파일 크기({file_size//1024//1024}MB)가 Groq 제한(24MB)을 초과합니다.\n"
                "파일을 분할하거나 압축 후 재시도해주세요."
            )
        gc = Groq(api_key=GROQ_API_KEY)
        with open(file_path, "rb") as f:
            result = gc.audio.transcriptions.create(
                file=(fname, f.read()),
                model=GROQ_WHISPER,
                language="ko",
                response_format="text",
            )
        text = result if isinstance(result, str) else getattr(result, "text", str(result))
        return text.strip()

    elif engine == "openai":
        from openai import OpenAI
        oc = OpenAI(api_key=OPENAI_API_KEY)
        with open(file_path, "rb") as f:
            result = oc.audio.transcriptions.create(
                model="whisper-1", file=f, language="ko")
        return result.text

    elif engine == "azure":
        with open(file_path, "rb") as f:
            result = azure_client.audio.transcriptions.create(
                model=AZURE_WHISPER, file=f)
        return result.text

    elif engine == "local":
        from faster_whisper import WhisperModel
        model = WhisperModel("base", device="cpu", compute_type="int8")
        segs, _ = model.transcribe(file_path, language="ko", beam_size=1, vad_filter=True)
        return " ".join(s.text.strip() for s in segs)

    raise ValueError(
        "음성 변환 엔진이 없습니다. GROQ_API_KEY 환경 변수를 설정해주세요.\n"
        "무료 키 발급: https://console.groq.com"
    )


# ── GPT-4o 회의록 생성 ────────────────────────────────────────────────────────
def make_minutes(transcript: str, meta_text: str) -> dict:
    if not azure_client:
        raise ValueError("Azure OpenAI 키가 설정되지 않았습니다.")

    resp = azure_client.chat.completions.create(
        model=AZURE_DEPLOY,
        messages=[
            {"role": "system", "content":
             "당신은 이노티움(주) 전문 회의록 작성자입니다.\n"
             "원문에 있는 내용만 포함하세요. 추측·없는 내용 추가 금지.\n"
             "결정사항은 '~하기로 함' 형식. 격식체 사용.\n"
             "모든 배열 요소는 반드시 문자열(string)이어야 합니다. 객체(object) 사용 금지."},
            {"role": "user", "content":
             f"[회의 정보]\n{meta_text}\n\n[회의 내용]\n{transcript}\n\n"
             "아래 JSON 스키마를 엄격히 따르세요. 모든 값은 문자열(string)입니다:\n"
             '{"title": "문자열", "overview": ["문자열1", "문자열2"], '
             '"discussion": [{"main": "문자열", "sub": ["문자열1", "문자열2"]}], '
             '"decisions": ["문자열1 (~하기로 함)"], "others": ["문자열"]}'},
        ],
        response_format={"type": "json_object"},
        temperature=0.2,
        max_tokens=2000,
    )
    raw = __import__("json").loads(resp.choices[0].message.content)
    # 혹시 dict가 섞여 들어온 경우 문자열로 변환
    def _flatten(v):
        if isinstance(v, str): return v
        if isinstance(v, dict): return " ".join(str(x) for x in v.values())
        return str(v)
    raw["overview"]   = [_flatten(x) for x in raw.get("overview", [])]
    raw["decisions"]  = [_flatten(x) for x in raw.get("decisions", [])]
    raw["others"]     = [_flatten(x) for x in raw.get("others", [])]
    raw["discussion"] = [
        {"main": _flatten(d.get("main", d) if isinstance(d, dict) else d),
         "sub": [_flatten(s) for s in (d.get("sub", []) if isinstance(d, dict) else [])]}
        for d in raw.get("discussion", [])
    ]
    return raw


# ── 미리보기 / DOCX ─────────────────────────────────────────────────────────
def build_preview(m, date_str, place, attendees, author):
    lines = ["=" * 54, "              회  의  록", "=" * 54,
             f"  회의명   : {m.get('title','')}", f"  작성자   : {author}",
             f"  일  시   : {date_str}", f"  장  소   : {place}",
             f"  참석자   : {attendees}", "=" * 54,
             "  회의 내용 (주요 안건 및 협의 사항)", "=" * 54,
             "", "  회의 결과 아래와 같이 논의되었음을 안내드립니다.",
             "", "              - 아  래 -", "", "□ 개  요"]
    for it in m.get("overview", []):
        lines.append(f"  ▪ {it}")
    lines += ["", "□ 논의사항"]
    for d in m.get("discussion", []):
        lines.append(f"  ■ {d.get('main','')}")
        for s in d.get("sub", []):
            lines.append(f"      • {s}")
    lines += ["", "□ 결정사항"]
    for dec in m.get("decisions", []):
        lines.append(f"  ■ {dec}")
    if m.get("others"):
        lines += ["", "□ 기타"]
        for ot in m.get("others", []):
            lines.append(f"  ■ {ot}")
    lines += ["", "=" * 54]
    return "\n".join(lines)


def build_docx(m, date_str, place, attendees, author, out_path):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def shade(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def cw(cell, text, bold=False, size=10, align=None):
        p = cell.paragraphs[0]; p.clear(); r = p.add_run(text)
        r.font.bold = bold; r.font.size = Pt(size)
        if align: p.alignment = align

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
    style = doc.styles["Normal"]; style.font.name = "맑은 고딕"; style.font.size = Pt(10)

    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(12)
    tr = tp.add_run("회  의  록"); tr.font.size = Pt(20); tr.font.bold = True

    tbl = doc.add_table(rows=4, cols=4); tbl.style = "Table Grid"
    widths = [Cm(2.8), Cm(7.5), Cm(2.8), Cm(5.2)]
    for ci, w in enumerate(widths):
        for row in tbl.rows: row.cells[ci].width = w

    r0 = tbl.rows[0]; cw(r0.cells[0], "회  의  명", bold=True); shade(r0.cells[0], "D9D9D9")
    cw(r0.cells[1], m.get("title", "")); cw(r0.cells[2], "작  성  자", bold=True); shade(r0.cells[2], "D9D9D9")
    cw(r0.cells[3], author)

    r1 = tbl.rows[1]; cw(r1.cells[0], "회  의  일  시", bold=True); shade(r1.cells[0], "D9D9D9")
    cw(r1.cells[1], date_str); cw(r1.cells[2], "회  의  장  소", bold=True); shade(r1.cells[2], "D9D9D9")
    cw(r1.cells[3], place)

    r2 = tbl.rows[2]; merged = r2.cells[1].merge(r2.cells[3])
    cw(r2.cells[0], "참  석  자", bold=True); shade(r2.cells[0], "D9D9D9"); cw(merged, attendees)

    r3 = tbl.rows[3]; full = r3.cells[0].merge(r3.cells[3])
    cw(full, "회의 내용 (주요 안건 및 협의 사항)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(full, "BFBFBF")

    doc.add_paragraph()
    doc.add_paragraph("회의 결과 아래와 같이 논의되었음을 안내드립니다.")
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run("- 아  래 -").font.size = Pt(10)

    def sec_head(title):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
        r = p.add_run(title); r.font.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    def bullet(text, indent=0.5, marker="▪"):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(2); r = p.add_run(f"{marker} {text}"); r.font.size = Pt(10)

    sec_head("□ 개  요")
    for it in m.get("overview", []): bullet(it)
    sec_head("□ 논의사항")
    for d in m.get("discussion", []):
        bullet(d.get("main", ""), marker="■"); [bullet(s, indent=1.3, marker="•") for s in d.get("sub", [])]
    sec_head("□ 결정사항")
    for dec in m.get("decisions", []): bullet(dec)
    if m.get("others"):
        sec_head("□ 기타")
        for ot in m.get("others", []): bullet(ot)

    doc.save(out_path)


# ── 문서 텍스트 추출 ──────────────────────────────────────────────────────────
def extract_text_from_file(file_path: str, suffix: str) -> str:
    if suffix in (".txt", ".md", ".csv", ".rtf"):
        raw = Path(file_path).read_bytes()
        try:
            import chardet; enc = chardet.detect(raw).get("encoding") or "utf-8"
        except ImportError:
            enc = "utf-8"
        for fb in (enc, "utf-8", "euc-kr", "cp949"):
            try: return raw.decode(fb)
            except: pass
        return raw.decode("utf-8", errors="replace")
    elif suffix == ".pdf":
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    elif suffix == ".docx":
        from docx import Document as D
        return "\n".join(p.text for p in D(file_path).paragraphs if p.text.strip())
    elif suffix == ".xlsx":
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        rows = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                rt = "\t".join(str(c) for c in row if c is not None)
                if rt.strip(): rows.append(rt)
        return "\n".join(rows)
    elif suffix == ".hwp":
        import olefile, zlib, re as _re
        ole = olefile.OleFileIO(file_path); texts = []
        for entry in ole.listdir():
            nm = "/".join(entry)
            if "BodyText" in nm and "Section" in nm:
                try:
                    data = zlib.decompress(ole.openstream(nm).read(), -15)
                    text = _re.sub(r"[\x00-\x08\x0b-\x1f\x7f\x80-\x9f]", " ",
                                   data.decode("utf-16-le", errors="replace"))
                    if text.strip(): texts.append(text)
                except: pass
        return "\n".join(texts)
    else:
        raise ValueError(f"지원하지 않는 형식: {suffix}")


# ── 레드마인 ──────────────────────────────────────────────────────────────────
def make_markdown(m, date_str, place, attendees, author):
    lines = [f"# {m.get('title','회의록')}", "",
             "| 항목 | 내용 |", "|------|------|",
             f"| 회의 일시 | {date_str} |", f"| 회의 장소 | {place} |",
             f"| 참석자 | {attendees} |", f"| 작성자 | {author} |",
             "", "## 개요"]
    for it in m.get("overview", []): lines.append(f"- {it}")
    lines += ["", "## 논의사항"]
    for d in m.get("discussion", []):
        lines.append(f"- **{d.get('main','')}**")
        for s in d.get("sub", []): lines.append(f"  - {s}")
    lines += ["", "## 결정사항"]
    for dec in m.get("decisions", []): lines.append(f"- {dec}")
    if m.get("others"):
        lines += ["", "## 기타"]
        for ot in m.get("others", []): lines.append(f"- {ot}")
    lines += ["", "---", f"*자동 생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}*"]
    return "\n".join(lines)


# ── Flask 라우트 ──────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template_string(HTML)

@app.route("/health")
def health():
    engine, label = _stt_engine()
    return __import__("json").dumps(
        {"ok": True, "server": "meeting-minutes-cloud", "version": "3.0",
         "stt": label or "unavailable", "gpt4o": bool(AZURE_KEY)},
        ensure_ascii=False
    ), 200, {"Content-Type": "application/json; charset=utf-8"}

@app.route("/stt-status")
def stt_status():
    engine, label = _stt_engine()
    return __import__("json").dumps(
        {"available": engine is not None, "engine": label or "없음"},
        ensure_ascii=False
    ), 200, {"Content-Type": "application/json; charset=utf-8"}

@app.route("/generate", methods=["POST"])
def generate():
    if request.is_json:
        data = request.get_json(force=True) or {}
        text = data.get("text", "").strip()
        if not text:
            return __import__("json").dumps({"error": "텍스트를 입력해주세요."}), 400, {"Content-Type": "application/json"}
        date_str = data.get("date", "").replace("T", " ")
        place = data.get("place", ""); attendees = data.get("attendees", ""); author = data.get("author", "")
    else:
        mode = request.form.get("mode", "audio")
        date_str = request.form.get("date", "").replace("T", " ")
        place = request.form.get("place", ""); attendees = request.form.get("attendees", ""); author = request.form.get("author", "")
        if mode == "audio":
            if "audio" not in request.files:
                return __import__("json").dumps({"error": "음성 파일이 없습니다."}), 400, {"Content-Type": "application/json"}
            audio = request.files["audio"]
            suffix = Path(audio.filename).suffix.lower()
            if suffix not in ALLOWED_AUDIO:
                return __import__("json").dumps({"error": f"지원하지 않는 형식: {suffix}"}), 400, {"Content-Type": "application/json"}
            tmp = UPLOAD_TMP / f"rec_{datetime.now().strftime('%Y%m%d_%H%M%S')}{suffix}"
            audio.save(str(tmp))
            try:
                text = transcribe(str(tmp))
            except Exception as e:
                return __import__("json").dumps({"error": str(e)}, ensure_ascii=False), 500, {"Content-Type": "application/json; charset=utf-8"}
            finally:
                tmp.unlink(missing_ok=True)
            if not text or len(text.strip()) < 3:
                return __import__("json").dumps({
                    "error": "음성에서 텍스트를 인식하지 못했습니다.\n"
                             "① 녹음 볼륨이 충분한지 확인해주세요.\n"
                             "② 지원 형식: m4a · mp3 · wav · mp4 (최대 25MB)\n"
                             "③ 파일이 25MB 초과 시 분할 후 업로드하세요."
                }, ensure_ascii=False), 400, {"Content-Type": "application/json; charset=utf-8"}
        else:
            text = request.form.get("text", "").strip()

    meta = []
    if date_str: meta.append(f"- 일시: {date_str}")
    if place:    meta.append(f"- 장소: {place}")
    if attendees: meta.append(f"- 참석자: {attendees}")
    if author:   meta.append(f"- 작성자: {author}")
    meta_text = "\n".join(meta) or "- (입력 없음, 내용에서 자동 추론)"

    try:
        minutes = make_minutes(text, meta_text)
    except Exception as e:
        return __import__("json").dumps({"error": f"회의록 생성 실패: {e}"}, ensure_ascii=False), 500, {"Content-Type": "application/json; charset=utf-8"}

    preview = build_preview(minutes, date_str, place, attendees, author)
    filename = f"[회의록] {minutes.get('title', datetime.now().strftime('%Y%m%d'))}.docx"

    resp_data = __import__("json").dumps({
        "minutes": minutes, "preview": preview,
        "filename": filename, "date": date_str,
        "place": place, "attendees": attendees, "author": author,
    }, ensure_ascii=False)
    return resp_data, 200, {"Content-Type": "application/json; charset=utf-8"}

@app.route("/download", methods=["POST"])
def download():
    data = request.get_json()
    m = data.get("minutes", {}); filename = data.get("filename", "회의록.docx")
    out = UPLOAD_TMP / filename
    try:
        build_docx(m, data.get("date",""), data.get("place",""), data.get("attendees",""), data.get("author",""), str(out))
    except Exception as e:
        return __import__("json").dumps({"error": f"DOCX 생성 실패: {e}"}, ensure_ascii=False), 500, {"Content-Type": "application/json"}
    return send_file(str(out), as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/extract-text", methods=["POST"])
def extract_text_route():
    if "file" not in request.files:
        return __import__("json").dumps({"error": "파일이 없습니다."}), 400, {"Content-Type": "application/json"}
    f = request.files["file"]; suffix = Path(f.filename).suffix.lower()
    if suffix not in ALLOWED_DOCS:
        return __import__("json").dumps({"error": f"지원하지 않는 형식: {suffix}"}), 400, {"Content-Type": "application/json"}
    tmp = UPLOAD_TMP / f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}{suffix}"
    f.save(str(tmp))
    try:
        text = extract_text_from_file(str(tmp), suffix).strip()
        if not text:
            return __import__("json").dumps({"error": "텍스트 추출 실패"}), 400, {"Content-Type": "application/json"}
        return __import__("json").dumps({"text": text, "chars": len(text), "filename": f.filename}, ensure_ascii=False), 200, {"Content-Type": "application/json; charset=utf-8"}
    except Exception as e:
        return __import__("json").dumps({"error": str(e)}, ensure_ascii=False), 400, {"Content-Type": "application/json; charset=utf-8"}
    finally:
        tmp.unlink(missing_ok=True)

@app.route("/redmine", methods=["POST"])
def redmine():
    if not REDMINE_URL or not REDMINE_API_KEY:
        return __import__("json").dumps({"error": "레드마인이 설정되지 않았습니다. 사내망에서 접속 후 이용해주세요."}, ensure_ascii=False), 503, {"Content-Type": "application/json; charset=utf-8"}
    data = request.get_json()
    m = data.get("minutes", {}); date_str = data.get("date",""); place = data.get("place","")
    attendees = data.get("attendees",""); author = data.get("author","")
    title = m.get("title","회의록")
    date_only = date_str[:10] if date_str else datetime.now().strftime("%Y-%m-%d")
    subject = f"[회의록] {title} ({date_only})"
    content = make_markdown(m, date_str, place, attendees, author)
    hdrs = {"X-Redmine-API-Key": REDMINE_API_KEY, "Content-Type": "application/json"}
    body = __import__("json").dumps({
        "issue": {"project_id": REDMINE_PROJECT, "tracker_id": REDMINE_TRACKER_UPLOAD,
                  "status_id": 5, "subject": subject, "description": content}
    }, ensure_ascii=False).encode("utf-8")
    try:
        r = http_req.post(f"{REDMINE_URL}/issues.json", data=body, headers=hdrs, timeout=8)
        if r.status_code in (200, 201):
            issue_id = r.json()["issue"]["id"]
            return __import__("json").dumps({"ok": True, "url": f"{REDMINE_URL}/issues/{issue_id}"}, ensure_ascii=False), 200, {"Content-Type": "application/json; charset=utf-8"}
        return __import__("json").dumps({"error": f"레드마인 오류 (HTTP {r.status_code})"}, ensure_ascii=False), 500, {"Content-Type": "application/json; charset=utf-8"}
    except Exception as e:
        return __import__("json").dumps({"error": f"연결 실패: {e} — 사내망 연결을 확인하세요."}, ensure_ascii=False), 503, {"Content-Type": "application/json; charset=utf-8"}

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5002))
    app.run(host="0.0.0.0", port=port, debug=False)
